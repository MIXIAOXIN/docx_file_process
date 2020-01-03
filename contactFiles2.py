#coding=utf-8
import os
import pydoc
import docx
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docxcompose.composer import Composer



in_dir = "/Volumes/LACIE SHARE/7_其他文档/研究生课程/工程案例助教/7/7"
out_dir = "/Volumes/LACIE SHARE/7_其他文档/研究生课程/工程案例助教/7"

filenames=os.listdir(in_dir)

out_contact_file = os.path.join(out_dir, "seven.docx")
docx.Document().save(out_contact_file)
master = docx.Document(out_contact_file)
composer = Composer(master)

#f=open(out_contact_file, 'a+')


def chg_font(obj, fontname='微软雅黑', size=None):
    ## 设置字体函数
    obj.font.name = fontname
    obj._element.rPr.rFonts.set(qn('w:eastAsia'), fontname)
    if size and isinstance(size, Pt):
        obj.font.size = size

def combine_all_docx(filename_master,files_list):
    number_of_sections=len(files_list)
    master = docx.Document(filename_master)
    composer = Composer(master)
    for i in range(1, number_of_sections):
        doc_temp = docx.Document(files_list[i])
        composer.append(doc_temp)
    composer.save("combined_file.docx")


#先遍历文件名
file_number = 0
for filename in filenames:
    if filename[0] == '2':
        file_number += 1
        filepath = in_dir + '/'+filename
        print(filepath)

        ID = filename.split('_')
        NAME = ID[-1].split('.d')
        temp_name_docx = docx.Document()
        temp_name_docx.add_heading(u'分数：', 0)
        temp_name_docx.add_paragraph('姓名：' +NAME[0])
        temp_name_docx.add_paragraph('学号：' + ID[0])
        temp_name_docx.add_paragraph('\n')

        ##### 方法1： 按行遍历   （开始）#################################
        #遍历单个文件，读取行数
        # docx_file = docx.Document(filepath)
        # docx_paragraphs = docx_file.paragraphs
        # ps_detail = [(x.text, x.style.name) for x in docx_paragraphs]
        # for line in ps_detail:
        #     my_doc.add_paragraph(line[0])
            #f.write(line[0] + '\t' + line[1] + '\n\n')
        ##### 方法1： 按行遍历   （结束）#################################
        #f.write('\n')

        ##### 方法2： 文档复制   （开始）#################################
        #docx_name = Document_compose(temp_name_docx)
        print('1')
        docx_temp = docx.Document(filepath)
        chg_font(docx_temp.styles['Normal'], fontname='宋体', size=Pt(12))
        composer.append(temp_name_docx)
        print('2')
        composer.append(docx_temp)
        print('3')

        ##### 方法2： 文档复制   （结束）#################################
        temp_page_docx = docx.Document()
        temp_page_docx.add_page_break()
        composer.append(temp_page_docx)
#关闭文件

print('contact files finished.')
print('一共提交作业：', file_number)
#chg_font(my_doc.styles['Normal'], fontname='宋体', size=Pt(11))
#my_doc.save(out_contact_file)

composer.save(out_contact_file)
new_docx = docx.Document(out_contact_file)
chg_font(new_docx.styles['Normal'], fontname='宋体', size=Pt(12))
new_docx.save(out_contact_file)


