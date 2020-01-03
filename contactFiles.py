#coding=utf-8
import os
import pydoc
import docx
from docx.shared import Inches, Pt
from docx.oxml.ns import qn


in_dir = "/Volumes/LACIE SHARE/7_其他文档/研究生课程/工程案例助教/6/课程作业6"
out_dir = "/Volumes/LACIE SHARE/7_其他文档/研究生课程/工程案例助教/6"

filenames=os.listdir(in_dir)

out_contact_file = os.path.join(out_dir, "第6次作业汇总.docx")
my_doc = docx.Document()

# f=open(out_contact_file, 'a+')


def chg_font(obj, fontname='微软雅黑', size=None):
    ## 设置字体函数
    obj.font.name = fontname
    obj._element.rPr.rFonts.set(qn('w:eastAsia'), fontname)
    if size and isinstance(size, Pt):
        obj.font.size = size




#先遍历文件名
file_number = 0
for filename in filenames:
    if filename[0] == '2':
        file_number += 1
        filepath = in_dir + '/'+filename
        print(filepath)

        ID = filename.split('_')
        NAME = ID[-1].split('.d')
        my_doc.add_heading(u'分数：', 0)
        # my_doc.add_paragraph('分数：')
        my_doc.add_paragraph('姓名：' +NAME[0])
        my_doc.add_paragraph('学号：' + ID[0])
        my_doc.add_paragraph('\n')
        #遍历单个文件，读取行数
        docx_file = docx.Document(filepath)
        docx_paragraphs = docx_file.paragraphs
        ps_detail = [(x.text, x.style.name) for x in docx_paragraphs]
        for line in ps_detail:
            my_doc.add_paragraph(line[0])
            #f.write(line[0] + '\t' + line[1] + '\n\n')

        #f.write('\n')
        my_doc.add_page_break()
#关闭文件

print('contact files finished.')
print('一共提交作业：', file_number)
#f.close()
chg_font(my_doc.styles['Normal'], fontname='宋体', size=Pt(11))
my_doc.save(out_contact_file)

